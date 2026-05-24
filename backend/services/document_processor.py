import uuid
import re
from pathlib import Path
from typing import List
from datetime import datetime

from models.schemas import DocumentChunk, DocumentMetadata
from config import settings
from utils.logger import get_logger

logger = get_logger(__name__)


class DocumentProcessor:

    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap

    def process_file(self, file_path: Path, filename: str) -> tuple[DocumentMetadata, List[DocumentChunk]]:
        document_id = str(uuid.uuid4())
        file_type = file_path.suffix.lower()
        file_size_kb = round(file_path.stat().st_size / 1024, 2)

        logger.info(f"Processing '{filename}' — {file_size_kb}KB, type: {file_type}")

        pages = self._extract_text(file_path, file_type)
        chunks = self._create_chunks(pages, document_id, filename)

        logger.info(f"Created {len(chunks)} chunks from '{filename}'")

        metadata = DocumentMetadata(
            document_id=document_id,
            filename=filename,
            file_type=file_type,
            total_chunks=len(chunks),
            upload_time=datetime.now(),
            file_size_kb=file_size_kb,
        )

        return metadata, chunks

    def _extract_text(self, file_path: Path, file_type: str) -> List[dict]:
        if file_type == ".pdf":
            return self._extract_pdf(file_path)
        elif file_type == ".docx":
            return self._extract_docx(file_path)
        elif file_type == ".txt":
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf(self, file_path: Path) -> List[dict]:
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(file_path))
            if reader.is_encrypted:
                raise ValueError("PDF is password-protected. Please remove the password and re-upload.")
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append({"page_number": i + 1, "text": self._clean_text(text)})
            if not pages:
                logger.warning(f"No text extracted from PDF '{file_path.name}' — may be image-only")
                raise ValueError(
                    "No text could be extracted from this PDF. "
                    "It may be a scanned document or contain only images. "
                    "Try running it through an OCR tool first."
                )
            logger.info(f"Extracted {len(pages)} pages from '{file_path.name}'")
            return pages
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"PDF extraction failed for '{file_path.name}': {str(e)}")
            raise ValueError(
                "Could not read this PDF. It may be corrupted, not a real PDF file, "
                "or contain only scanned images. Try opening it in Preview first to confirm it works."
            )

    def _extract_docx(self, file_path: Path) -> List[dict]:
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
            try:
                doc = Document(str(file_path))
            except PackageNotFoundError:
                raise ValueError("This DOCX file appears to be corrupted or is not a valid Word document.")

            parts = []
            for p in doc.paragraphs:
                if p and p.text and p.text.strip():
                    parts.append(p.text.strip())

            for table in (doc.tables or []):
                if not table:
                    continue
                for row in (table.rows or []):
                    if not row:
                        continue
                    seen = []
                    for cell in (row.cells or []):
                        if not cell:
                            continue
                        text = cell.text.strip() if cell.text else ""
                        if text and text not in seen:
                            seen.append(text)
                    if seen:
                        parts.append(' | '.join(seen))

            if not parts:
                logger.warning(f"No text extracted from DOCX '{file_path.name}'")
                raise ValueError(
                    "No text could be extracted from this Word document. "
                    "It may be empty or contain only images."
                )
            logger.info(f"Extracted {len(parts)} text blocks from '{file_path.name}'")
            full_text = "\n".join(parts)
            return [{"page_number": 1, "text": self._clean_text(full_text)}]
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed for '{file_path.name}': {str(e)}")
            raise ValueError(f"Could not read DOCX: {str(e)}")

    def _extract_txt(self, file_path: Path) -> List[dict]:
        try:
            import chardet                         
            raw = file_path.read_bytes()
            if not raw.strip():
                raise ValueError("This text file is empty.")
            detected = chardet.detect(raw)        
            encoding = detected.get("encoding") or "utf-8"
            confidence = detected.get("confidence") or 0
            if confidence < 0.5:
                encoding = "utf-8"
            try:
                text = raw.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                text = raw.decode("utf-8", errors="replace")
            if not text.strip():
                raise ValueError("This text file contains no readable text.")
            logger.info(f"Extracted text from '{file_path.name}' using encoding '{encoding}'")
            return [{"page_number": 1, "text": self._clean_text(text)}]
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"TXT extraction failed for '{file_path.name}': {str(e)}")
            raise ValueError(f"Could not read TXT file: {str(e)}")

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def _create_chunks(self, pages: List[dict], document_id: str, filename: str) -> List[DocumentChunk]:
        all_chunks = []
        chunk_index = 0

        for page in pages:
            text = page["text"]
            page_number = page["page_number"]

            start = 0
            while start < len(text):
                end = start + self.chunk_size
                chunk_text = text[start:end]

                if len(chunk_text.strip()) < 50:
                    break

                all_chunks.append(DocumentChunk(
                    chunk_id=str(uuid.uuid4()),
                    document_id=document_id,
                    filename=filename,
                    content=chunk_text.strip(),
                    page_number=page_number,
                    chunk_index=chunk_index,
                    total_chunks=0,
                ))

                chunk_index += 1
                start += self.chunk_size - self.chunk_overlap

        total = len(all_chunks)
        for chunk in all_chunks:
            chunk.total_chunks = total

        return all_chunks

    def extract_full_text(self, file_path: Path) -> str:
        """Returns complete document text without chunking — used for preview."""
        file_type = file_path.suffix.lower()
        pages = self._extract_text(file_path, file_type)
        return "\n\n".join(page["text"] for page in pages)